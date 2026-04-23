from __future__ import annotations

import argparse
import base64
import json
import os
import re
import sys
import tempfile
import time
from dataclasses import dataclass
from pathlib import Path
from typing import Any
from urllib.parse import parse_qs, urlparse

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from selenium import webdriver
from selenium.webdriver.chrome.options import Options as ChromeOptions
from selenium.webdriver.common.by import By
from selenium.webdriver.edge.options import Options as EdgeOptions


XTBZ_HEADER = "ykt"


@dataclass
class Settings:
    browser: str
    browser_binary: str | None
    user_data_dir: str
    profile_directory: str
    course_url: str
    classroom_id: str
    output_dir: Path
    document_title: str
    docx_name: str
    save_raw: bool
    save_images: bool
    include_source_url: bool
    limit_homeworks: int | None
    headless: bool
    startup_wait: float


def infer_classroom_id(course_url: str) -> str | None:
    parsed = urlparse(course_url)
    query = parse_qs(parsed.query)
    for key in ("classroom_id", "cid"):
        values = query.get(key)
        if values and values[0]:
            return values[0]

    match = re.search(r"/studentLog/(\d+)", parsed.path)
    if match:
        return match.group(1)
    return None


def default_user_data_dir(browser: str) -> str | None:
    local_appdata = os.environ.get("LOCALAPPDATA")
    if not local_appdata:
        return None

    if browser == "edge":
        return str(Path(local_appdata) / "Microsoft/Edge/User Data")
    if browser == "chrome":
        return str(Path(local_appdata) / "Google/Chrome/User Data")
    return None


def normalize_docx_name(name: str) -> str:
    return name if name.lower().endswith(".docx") else f"{name}.docx"


def parse_args() -> Settings:
    parser = argparse.ArgumentParser(
        description="导出雨课堂作业到 Word。默认复用本机已登录浏览器资料。"
    )
    parser.add_argument(
        "--course-url",
        required=True,
        help="已登录后可以打开的雨课堂课程页面 URL，例如 studentLog 页面。",
    )
    parser.add_argument(
        "--classroom-id",
        help="可选。若未提供，则从 --course-url 中自动推断 classroom_id。",
    )
    parser.add_argument(
        "--browser",
        choices=("edge", "chrome"),
        default="edge",
        help="要复用登录态的浏览器。默认 edge。",
    )
    parser.add_argument(
        "--browser-binary",
        help="浏览器可执行文件路径。默认使用系统已安装浏览器。",
    )
    parser.add_argument(
        "--user-data-dir",
        help="浏览器 User Data 目录。Windows 下默认自动推断。",
    )
    parser.add_argument(
        "--profile-directory",
        default="Default",
        help="浏览器配置目录名，例如 Default、Profile 1。默认 Default。",
    )
    parser.add_argument(
        "--output-dir",
        default="output",
        help="输出目录。默认 ./output",
    )
    parser.add_argument(
        "--document-title",
        default="雨课堂作业汇编",
        help="Word 文档标题。默认“雨课堂作业汇编”。",
    )
    parser.add_argument(
        "--docx-name",
        default="yuketang_homework_export.docx",
        help="输出 Word 文件名。默认 yuketang_homework_export.docx",
    )
    parser.add_argument(
        "--limit-homeworks",
        type=int,
        help="仅处理前 N 份作业，用于调试。",
    )
    parser.add_argument(
        "--save-raw",
        action="store_true",
        help="保存原始接口 JSON 到 output/raw_json。",
    )
    parser.add_argument(
        "--save-images",
        action="store_true",
        help="保留题目截图到 output/images。默认只把截图嵌入 Word，不额外保留图片文件。",
    )
    parser.add_argument(
        "--include-source-url",
        action="store_true",
        help="在 Word 首页写入课程来源 URL。",
    )
    parser.add_argument(
        "--no-headless",
        action="store_true",
        help="关闭无头模式，调试时可查看浏览器窗口。",
    )
    parser.add_argument(
        "--startup-wait",
        type=float,
        default=5.0,
        help="打开课程页后的等待秒数。默认 5 秒。",
    )

    args = parser.parse_args()

    classroom_id = args.classroom_id or infer_classroom_id(args.course_url)
    if not classroom_id:
        parser.error("无法从 --course-url 推断 classroom_id，请显式传入 --classroom-id。")

    user_data_dir = args.user_data_dir or default_user_data_dir(args.browser)
    if not user_data_dir:
        parser.error(
            "无法自动推断浏览器 User Data 目录，请显式传入 --user-data-dir。"
        )

    return Settings(
        browser=args.browser,
        browser_binary=args.browser_binary,
        user_data_dir=user_data_dir,
        profile_directory=args.profile_directory,
        course_url=args.course_url,
        classroom_id=classroom_id,
        output_dir=Path(args.output_dir).resolve(),
        document_title=args.document_title,
        docx_name=normalize_docx_name(args.docx_name),
        save_raw=args.save_raw,
        save_images=args.save_images,
        include_source_url=args.include_source_url,
        limit_homeworks=args.limit_homeworks,
        headless=not args.no_headless,
        startup_wait=args.startup_wait,
    )


def safe_name(value: str) -> str:
    value = re.sub(r'[\\\\/:*?"<>|]+', "_", value.strip())
    value = re.sub(r"\s+", " ", value)
    return value or "untitled"


def build_browser_options(settings: Settings):
    if settings.browser == "edge":
        options = EdgeOptions()
    else:
        options = ChromeOptions()

    if settings.browser_binary:
        options.binary_location = settings.browser_binary

    if settings.headless:
        options.add_argument("--headless=new")
    options.add_argument("--disable-gpu")
    options.add_argument("--disable-dev-shm-usage")
    options.add_argument("--window-size=1500,1200")
    options.add_argument("--log-level=3")
    options.add_argument(f"--user-data-dir={settings.user_data_dir}")
    options.add_argument(f"--profile-directory={settings.profile_directory}")
    return options


def open_logged_in_driver(settings: Settings):
    options = build_browser_options(settings)
    driver = webdriver.Edge(options=options) if settings.browser == "edge" else webdriver.Chrome(options=options)
    driver.get(settings.course_url)
    time.sleep(settings.startup_wait)
    return driver


def fetch_json(driver, path: str) -> dict[str, Any]:
    data = driver.execute_async_script(
        """
        const [path, xtbzHeader, done] = arguments;
        fetch(path, {
          credentials: "include",
          headers: { XTBZ: xtbzHeader || "ykt" }
        })
          .then(async (response) => {
            const text = await response.text();
            try {
              done({ ok: response.ok, status: response.status, payload: JSON.parse(text) });
            } catch (error) {
              done({ ok: false, status: response.status, raw: text, parseError: String(error) });
            }
          })
          .catch((error) => done({ ok: false, error: String(error) }));
        """,
        path,
        XTBZ_HEADER,
    )
    if not data.get("ok"):
        raise RuntimeError(f"接口请求失败: {path}\n{json.dumps(data, ensure_ascii=False)}")

    payload = data["payload"]
    if payload.get("success") is False:
        raise RuntimeError(f"接口返回失败: {path}\n{json.dumps(payload, ensure_ascii=False)}")
    return payload


def chapter_api_path(classroom_id: str) -> str:
    return f"/mooc-api/v1/lms/learn/course/chapter?cid={classroom_id}"


def collect_homeworks(course_chapter: list[dict[str, Any]]) -> list[dict[str, Any]]:
    homeworks: list[dict[str, Any]] = []
    for chapter_index, chapter in enumerate(course_chapter):
        chapter_name = chapter.get("name", f"Chapter {chapter_index + 1}")
        for section_index, item in enumerate(chapter.get("section_leaf_list", [])):
            if item.get("leaf_type") != 6:
                continue
            homeworks.append(
                {
                    "chapter_index": chapter_index,
                    "section_index": section_index,
                    "chapter_name": chapter_name,
                    "leaf_id": item["id"],
                    "name": item.get("name", f"{chapter_name} 作业"),
                    "leaf_order": item.get("order", section_index),
                }
            )
    return homeworks


def decode_base64_text(value: str) -> str:
    return base64.b64decode(value).decode("utf-8")


def load_homework_bundle(driver, classroom_id: str, leaf_id: int) -> dict[str, Any]:
    result = driver.execute_async_script(
        """
        const [leafId, classroomId, done] = arguments;

        const encodeText = (text) => {
          const bytes = new TextEncoder().encode(text);
          let binary = "";
          bytes.forEach((value) => { binary += String.fromCharCode(value); });
          return btoa(binary);
        };

        const fetchPayload = async (path) => {
          const response = await fetch(path, {
            credentials: "include",
            headers: { XTBZ: "ykt" },
          });
          const text = await response.text();
          return {
            ok: response.ok,
            status: response.status,
            text,
            payload: JSON.parse(text),
          };
        };

        (async () => {
          const leafResponse = await fetchPayload(`/mooc-api/v1/lms/learn/leaf_info/${classroomId}/${leafId}/?cid=${classroomId}`);
          const exerciseId = leafResponse.payload.data.content_info.leaf_type_id;
          const exerciseResponse = await fetchPayload(`/mooc-api/v1/lms/exercise/get_exercise_list/${exerciseId}/?cid=${classroomId}`);
          const exerciseData = exerciseResponse.payload.data;
          window.__codexExerciseData = exerciseData;
          window.__codexExerciseLeafId = leafId;
          done({
            ok: true,
            exercise_id: exerciseId,
            leaf_json_b64: encodeText(leafResponse.text),
            exercise_json_b64: encodeText(exerciseResponse.text),
            summary: {
              name: exerciseData.name,
              font_url: exerciseData.font || "",
              problems: exerciseData.problems.map((problem) => ({
                index: problem.index,
                score: problem.score,
                content: {
                  Type: problem.content.Type,
                  TypeText: problem.content.TypeText,
                  ProblemType: problem.content.ProblemType,
                },
                user: {
                  is_right: problem.user && problem.user.is_right,
                  my_answer: problem.user && problem.user.my_answer ? problem.user.my_answer : null,
                  my_answers: problem.user && problem.user.my_answers ? problem.user.my_answers : null,
                  my_score: problem.user && problem.user.my_score ? problem.user.my_score : "0",
                  submit_time: problem.user && problem.user.submit_time ? problem.user.submit_time : "",
                },
              })),
            },
          });
        })().catch((error) => done({ ok: false, error: String(error) }));
        """,
        leaf_id,
        classroom_id,
    )
    if not result.get("ok"):
        raise RuntimeError(f"作业抓取失败: {leaf_id}\n{json.dumps(result, ensure_ascii=False)}")
    return result


def render_problem_image(driver, problem_index: int, image_path: Path) -> None:
    result = driver.execute_async_script(
        """
        const [problemIndex, done] = arguments;
        const data = window.__codexExerciseData;
        if (!data || !Array.isArray(data.problems)) {
          done({ ok: false, error: "window.__codexExerciseData 不存在" });
          return;
        }
        const problem = data.problems[problemIndex];
        if (!problem) {
          done({ ok: false, error: `未找到题目索引 ${problemIndex}` });
          return;
        }

        const fontUrl = data.font || "";
        let style = document.getElementById("codex-problem-style");
        if (!style) {
          style = document.createElement("style");
          style.id = "codex-problem-style";
          document.head.appendChild(style);
        }
        style.textContent = `
          @font-face { font-family: "exam-data-decrypt-font"; src: url("${fontUrl}"); }
          html, body { margin: 0; padding: 0; background: #ffffff; }
          body { padding: 24px; font-family: "Microsoft YaHei", "PingFang SC", sans-serif; color: #0f172a; }
          #codex-problem-card { width: 1120px; padding: 26px 30px; border: 1px solid #dbe3ef; border-radius: 18px; background: #ffffff; }
          #codex-problem-card p, #codex-problem-card span, #codex-problem-card div, #codex-problem-card li {
            font-size: 30px; line-height: 1.62;
          }
          #codex-problem-card p { margin: 0 0 12px; }
          #codex-problem-card .options { margin-top: 14px; display: grid; gap: 10px; }
          #codex-problem-card .option-row { display: flex; align-items: flex-start; gap: 12px; }
          #codex-problem-card .option-key { width: 48px; flex: 0 0 48px; font-weight: 700; color: #1e3a8a; }
          #codex-problem-card .option-content { flex: 1; }
          #codex-problem-card .blank-item {
            display: inline-block; min-width: 88px; padding: 0 10px; margin: 0 4px;
            border: 2px solid #94a3b8; border-radius: 8px; background: #f8fafc; text-align: center;
          }
          #codex-problem-card .xuetangx-com-encrypted-font {
            font-family: "exam-data-decrypt-font" !important;
          }
          #codex-problem-card img { max-width: 100%; height: auto; }
          #codex-problem-card table { border-collapse: collapse; }
          #codex-problem-card td, #codex-problem-card th { border: 1px solid #cbd5e1; padding: 6px 10px; }
        `;

        document.body.innerHTML = '<div id="codex-problem-card"><div class="problem-body"></div><div class="options"></div></div>';
        const card = document.getElementById("codex-problem-card");
        card.querySelector(".problem-body").innerHTML = problem.content.Body || "";

        const optionsWrap = card.querySelector(".options");
        const options = problem.content.Options || [];
        optionsWrap.innerHTML = options.map((option) => `
          <div class="option-row">
            <div class="option-key">${option.key || ""}.</div>
            <div class="option-content">${option.value || ""}</div>
          </div>
        `).join("");
        if (!options.length) {
          optionsWrap.remove();
        }

        const imagePromises = Array.from(card.querySelectorAll("img")).map((image) => {
          if (image.complete) return Promise.resolve();
          return new Promise((resolve) => {
            image.onload = resolve;
            image.onerror = resolve;
          });
        });

        Promise.all([
          document.fonts ? document.fonts.ready : Promise.resolve(),
          Promise.all(imagePromises),
          new Promise((resolve) => setTimeout(resolve, 300)),
        ]).then(() => {
          const rect = card.getBoundingClientRect();
          done({
            ok: true,
            width: Math.ceil(rect.width + 80),
            height: Math.ceil(rect.height + 80),
          });
        }).catch((error) => done({ ok: false, error: String(error) }));
        """,
        problem_index,
    )
    if not result.get("ok"):
        raise RuntimeError(f"题目渲染失败: {json.dumps(result, ensure_ascii=False)}")

    driver.set_window_size(max(result["width"], 1280), max(result["height"], 700))
    time.sleep(0.2)
    image_path.parent.mkdir(parents=True, exist_ok=True)
    driver.find_element(By.ID, "codex-problem-card").screenshot(str(image_path))


def answer_from_problem(problem: dict[str, Any]) -> str:
    content = problem.get("content", {})
    user = problem.get("user", {})
    problem_type = content.get("Type")

    if problem_type in {"SingleChoice", "Judgement"}:
        values = user.get("my_answer") or []
        if problem_type == "Judgement":
            mapping = {"true": "正确", "false": "错误", "True": "正确", "False": "错误"}
            return "、".join(mapping.get(str(value), str(value)) for value in values) or "未获取到"
        return "、".join(str(value) for value in values) or "未获取到"

    if problem_type == "MultipleChoice":
        values = user.get("my_answers") or {}
        selected = [key for key, enabled in sorted(values.items()) if enabled]
        return "、".join(selected) or "未获取到"

    if problem_type == "FillBlank":
        values = user.get("my_answers") or {}
        parts = []
        for key in sorted(values, key=lambda item: int(item) if str(item).isdigit() else str(item)):
            answer = values[key]
            if isinstance(answer, dict):
                parts.append(f"第{key}空：{answer.get('answer', '')}")
            else:
                parts.append(f"第{key}空：{answer}")
        return "；".join(parts) or "未获取到"

    if user.get("my_answer"):
        return json.dumps(user["my_answer"], ensure_ascii=False)
    if user.get("my_answers"):
        return json.dumps(user["my_answers"], ensure_ascii=False)
    return "未获取到"


def result_text(problem: dict[str, Any]) -> str:
    is_right = problem.get("user", {}).get("is_right")
    if is_right is True:
        return "正确"
    if is_right is False:
        return "错误"
    return "未知"


def set_run_font(run) -> None:
    run.font.name = "Microsoft YaHei"
    run.font.size = Pt(11)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")


def add_text_paragraph(document: Document, text: str) -> None:
    paragraph = document.add_paragraph()
    run = paragraph.add_run(text)
    set_run_font(run)


def add_homework_to_doc(
    document: Document,
    homework: dict[str, Any],
    image_paths: list[Path],
    show_chapter_heading: bool,
) -> None:
    if show_chapter_heading:
        document.add_heading(homework["chapter_name"], level=1)
    document.add_heading(homework["name"], level=2)

    summary = (
        f"题目数：{len(homework['problems'])}    "
        f"总分：{homework['total_score']:.2f}    "
        f"实得：{homework['earned_score']:.2f}"
    )
    add_text_paragraph(document, summary)

    for index, (problem, image_path) in enumerate(zip(homework["problems"], image_paths), start=1):
        title = f"题目 {index}（{problem['content'].get('TypeText', '未知题型')}）"
        add_text_paragraph(document, title)
        document.add_picture(str(image_path), width=Inches(5.9))

        answer_text = answer_from_problem(problem)
        score_text = f"{problem.get('user', {}).get('my_score', '0')} / {problem.get('score', 0)}"
        add_text_paragraph(document, "答案来源：我的作答")
        add_text_paragraph(document, f"作答内容：{answer_text}")
        add_text_paragraph(document, f"平台判定：{result_text(problem)}    得分：{score_text}")


def prepare_document(settings: Settings) -> Document:
    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)
    style = document.styles["Normal"]
    style.font.name = "Microsoft YaHei"
    style.font.size = Pt(11)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")

    document.add_heading(settings.document_title, level=0)
    if settings.include_source_url:
        add_text_paragraph(document, f"来源页面：{settings.course_url}")
    add_text_paragraph(
        document,
        "说明：题目截图按网页原始加密字体渲染；官方答案未开放时，采用你的作答结果。",
    )
    return document


def save_raw_bundle(raw_dir: Path, raw_name: str, bundle: dict[str, Any]) -> None:
    raw_dir.mkdir(parents=True, exist_ok=True)
    (raw_dir / f"{raw_name}_leaf.json").write_text(
        decode_base64_text(bundle["leaf_json_b64"]),
        encoding="utf-8",
    )
    (raw_dir / f"{raw_name}_exercise.json").write_text(
        decode_base64_text(bundle["exercise_json_b64"]),
        encoding="utf-8",
    )


def main() -> int:
    settings = parse_args()
    settings.output_dir.mkdir(parents=True, exist_ok=True)

    raw_dir = settings.output_dir / "raw_json"
    image_dir = settings.output_dir / "images"

    temp_images: tempfile.TemporaryDirectory[str] | None = None
    working_image_dir: Path
    if settings.save_images:
        image_dir.mkdir(parents=True, exist_ok=True)
        working_image_dir = image_dir
    else:
        temp_images = tempfile.TemporaryDirectory(prefix="yuketang_images_")
        working_image_dir = Path(temp_images.name)

    driver = None
    homeworks_data: list[dict[str, Any]] = []

    try:
        print("打开浏览器登录态并抓取课程数据...")
        driver = open_logged_in_driver(settings)
        chapter_payload = fetch_json(driver, chapter_api_path(settings.classroom_id))

        if settings.save_raw:
            raw_dir.mkdir(parents=True, exist_ok=True)
            (raw_dir / "course_chapter.json").write_text(
                json.dumps(chapter_payload, ensure_ascii=False, indent=2),
                encoding="utf-8",
            )

        homeworks = collect_homeworks(chapter_payload["data"]["course_chapter"])
        if settings.limit_homeworks:
            homeworks = homeworks[: settings.limit_homeworks]
        print(f"找到 {len(homeworks)} 份作业。")

        for idx, homework in enumerate(homeworks, start=1):
            print(f"[{idx}/{len(homeworks)}] 抓取 {homework['name']} ...")
            bundle = load_homework_bundle(driver, settings.classroom_id, homework["leaf_id"])
            exercise_id = bundle["exercise_id"]
            exercise_summary = bundle["summary"]

            if settings.save_raw:
                raw_name = safe_name(f"{homework['chapter_name']}_{homework['name']}")
                save_raw_bundle(raw_dir, raw_name, bundle)

            problems = exercise_summary["problems"]
            total_score = sum(float(problem.get("score", 0) or 0) for problem in problems)
            earned_score = sum(
                float(problem.get("user", {}).get("my_score", 0) or 0)
                for problem in problems
            )
            homeworks_data.append(
                {
                    **homework,
                    "exercise_id": exercise_id,
                    "problems": problems,
                    "total_score": total_score,
                    "earned_score": earned_score,
                }
            )

        if settings.save_raw:
            (raw_dir / "summary.json").write_text(
                json.dumps(
                    [
                        {
                            "chapter_name": item["chapter_name"],
                            "name": item["name"],
                            "leaf_id": item["leaf_id"],
                            "exercise_id": item["exercise_id"],
                            "problem_count": len(item["problems"]),
                            "total_score": item["total_score"],
                            "earned_score": item["earned_score"],
                        }
                        for item in homeworks_data
                    ],
                    ensure_ascii=False,
                    indent=2,
                ),
                encoding="utf-8",
            )

        print("开始渲染题目截图并生成 Word ...")
        document = prepare_document(settings)
        previous_chapter = None

        for homework_index, homework in enumerate(homeworks_data, start=1):
            load_homework_bundle(driver, settings.classroom_id, homework["leaf_id"])
            homework_dir = working_image_dir / safe_name(f"{homework_index:02d}_{homework['name']}")
            homework_dir.mkdir(parents=True, exist_ok=True)
            image_paths: list[Path] = []

            for problem_pos, problem in enumerate(homework["problems"]):
                image_path = homework_dir / f"{problem['index']:02d}.png"
                render_problem_image(driver, problem_pos, image_path)
                image_paths.append(image_path)

            add_homework_to_doc(
                document,
                homework,
                image_paths,
                show_chapter_heading=homework["chapter_name"] != previous_chapter,
            )
            previous_chapter = homework["chapter_name"]
            if homework_index != len(homeworks_data):
                document.add_page_break()

        docx_path = settings.output_dir / settings.docx_name
        document.save(docx_path)

        print(f"已生成 Word：{docx_path}")
        if settings.save_raw:
            print(f"原始 JSON：{raw_dir}")
        if settings.save_images:
            print(f"题目截图：{image_dir}")
        else:
            print("题目截图仅用于嵌入 Word，未额外保留到磁盘。")
        return 0

    finally:
        if driver is not None:
            driver.quit()
        if temp_images is not None:
            temp_images.cleanup()


if __name__ == "__main__":
    sys.exit(main())
